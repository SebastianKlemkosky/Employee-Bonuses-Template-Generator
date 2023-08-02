import locale
import os
from docx import Document
from docx.shared import Inches
from PIL import Image

# Mapping of column names to template locations
COLUMN_MAPPING = {
    'Employee Name': 'Employee_Name',
    'Business Entity': 'Business_Entity',
    'Job Title': 'Job_Title',
    'Department': 'Department',
    'Budget Area': 'Budget_Area',
    'AIP Tier': 'AIP_Tier',
    'AIP Type': 'AIP_Type',
    'Currency Symbol': 'CC',
    'Eligible Salary': 'Eligible_Salary',
    'Minimum Bonus': 'Minimum_Bonus',
    'CenterPoint (CP)': 'CenterPoint',
    'Total Payout': 'Total_Payout',

    'Type': 'Type',

    'T1 Name': 'T1_Metric_Target',
    'T1 Weight': 'T1_Weight',
    'T1 Achievement': 'T1_Achievement',
    'T1 Multiplier': 'T1_Multiplier',
    'T1 Payout': 'T1_Payout',
    'T1 Matrix Number (Bottom)': 'T1_Matrix_Number(Bottom)',

    'T2 Name': 'T2_Metric_Target',
    'T2 Weight': 'T2_Weight',
    'T2 Achievement': 'T2_Achievement',
    'T2 Multiplier': 'T2_Multiplier',
    'T2 Payout': 'T2_Payout',
    'T2 Matrix Number (Bottom)': 'T2_Matrix_Number(Bottom)',

    'T3 Name': 'T3_Metric_Target',
    'T3 Weight': 'T3_Weight',
    'T3 Achievement': 'T3_Achievement',
    'T3 Multiplier': 'T3_Multiplier',
    'T3 Payout': 'T3_Payout',
    'T3 Matrix Number (Bottom)': 'T3_Matrix_Number(Bottom)',

    'T4 Name': 'T4_Metric_Target',
    'T4 Weight': 'T4_Weight',
    'T4 Achievement': 'T4_Achievement',
    'T4 Multiplier': 'T4_Multiplier',
    'T4 Payout': 'T4_Payout',
    'T4 Matrix Number (Bottom)': 'T4_Matrix_Number(Bottom)',

    'Number of Targets': 'Number_of_Targets',
    'Native $': 'Native_$'
}


def populate_images(doc, matrix_folder, template_location, value):
    # Check if the value is empty
    if not value:
        return doc

    # Construct the image path based on the matrix_folder and value
    image_path = os.path.join(matrix_folder, f"{value}.png")

    # Check if the image file exists
    if not os.path.isfile(image_path):
        return doc

    # Create a new paragraph in the document
    paragraph = doc.add_paragraph()

    # Center-align the paragraph
    paragraph.alignment = 1

    # Add the image to the paragraph
    run = paragraph.add_run()
    run.add_picture(image_path)

    # Replace the template_location with an empty string
    for paragraph in doc.paragraphs:
        if template_location in paragraph.text:
            paragraph.text = paragraph.text.replace(template_location, "")

    return doc


def populate_document(doc, template_location, value):
    for paragraph in doc.paragraphs:
        if template_location in paragraph.text:
            paragraph.text = paragraph.text.replace(template_location, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if template_location in cell.text:
                    if value:
                        cell.text = cell.text.replace(template_location, value)
                    else:
                        cell.text = cell.text.replace(template_location, "N/A")

    return doc


def remove_na_rows_from_table(doc):
    for table in doc.tables:
        rows_to_delete = []  # List to store rows to be deleted
        for i, row in enumerate(table.rows):
            if any(cell.text.strip() == "N/A" for cell in row.cells):
                rows_to_delete.append(i)

        for idx in reversed(rows_to_delete):  # Reverse the order to safely remove rows
            table._element.remove(table.rows[idx]._element)

    return doc

def format_money_value(value):
    try:
        money_value = float(value)
        return '{:,.2f}'.format(money_value)
    except ValueError:
        return value

def populate_template(template_path, matrix_folder, csv_row):
    doc = Document(template_path)

    # Get the first word from the Employee Name
    employee_name = csv_row.get('Employee Name', '')
    first_name = employee_name.split()[0] if employee_name else ''

    # Replace the placeholder "First_Name_Optional" with the first name
    doc = populate_document(doc, '«First_Name_Optional»', first_name)

    # Populate the remaining template with data from the CSV row
    for col, template_location in COLUMN_MAPPING.items():
        # Get the value from the CSV row
        cell_value = csv_row.get(col, '')
        # Remove the extra characters from the template location
        template_location = '«' + template_location + '»'

        # Convert weight columns to percentage if applicable
        if col.endswith('Weight'):
            try:
                weight = float(cell_value)
                cell_value = f'{weight * 100:.2f}%'
            except ValueError:
                pass
        elif col in ['CenterPoint (CP)', 'Minimum Bonus']:
            try:
                bonus = float(cell_value)
                cell_value = f'{bonus:.2f}%'
            except ValueError:
                pass
        elif col in ['Eligible Salary', 'T1 Payout', 'T2 Payout', 'T3 Payout', 'T4 Payout', 'Total Payout']:
            cell_value = format_money_value(cell_value)

        # Populate the document with the images
        if col.startswith('T') and col.endswith(' Matrix Number (Bottom)'):
            doc = populate_images(doc, matrix_folder, template_location, cell_value)

        # Remove rows with "N/A" in tables
        doc = remove_na_rows_from_table(doc)

        # Populate the document with the value
        doc = populate_document(doc, template_location, cell_value)

    return doc
